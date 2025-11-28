"""
Flask在线文件翻译应用
"""
import os
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import config
from file_parser import FileParser
from translator import Translator

app = Flask(__name__)
app.config['SECRET_KEY'] = config.SECRET_KEY
app.config['UPLOAD_FOLDER'] = config.UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = config.MAX_FILE_SIZE

# 确保上传文件夹存在
os.makedirs(config.UPLOAD_FOLDER, exist_ok=True)

# 初始化翻译器
translator = Translator()

def allowed_file(filename):
    """检查文件类型是否允许"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in config.ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """首页"""
    return render_template('index.html', languages=config.LANGUAGES, ai_models=config.AI_MODELS)

@app.route('/upload', methods=['POST'])
def upload_file():
    """处理文件上传"""
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({'error': '没有上传文件'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': '不支持的文件类型'}), 400
        
        # 保存文件
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # 解析文件
        file_ext = filename.rsplit('.', 1)[1].lower()
        
        # 对于Word和PDF文档，使用格式化解析
        if file_ext in ['docx', 'doc']:
            html_content, paragraphs = FileParser.parse_docx_with_format(file_path)
            # 清理临时文件
            os.remove(file_path)
            
            return jsonify({
                'success': True,
                'content': paragraphs,
                'html_content': html_content,
                'filename': filename,
                'has_format': True,
                'file_type': 'word'
            })
        elif file_ext == 'pdf':
            html_content, paragraphs = FileParser.parse_pdf_with_format(file_path)
            # 清理临时文件
            os.remove(file_path)
            
            return jsonify({
                'success': True,
                'content': paragraphs,
                'html_content': html_content,
                'filename': filename,
                'has_format': True,
                'file_type': 'pdf'
            })
        else:
            # 其他文件类型使用普通解析
            parsed_content = FileParser.parse_file(file_path, file_ext)
            # 清理临时文件
            os.remove(file_path)
            
            return jsonify({
                'success': True,
                'content': parsed_content,
                'filename': filename,
                'has_format': False
            })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/translate', methods=['POST'])
def translate():
    """翻译文本"""
    try:
        data = request.get_json()
        
        if not data or 'content' not in data:
            return jsonify({'error': '缺少内容'}), 400
        
        content = data['content']
        target_lang = data.get('target_lang', 'zh-CN')
        source_lang = data.get('source_lang', 'auto')
        html_content = data.get('html_content')  # 原始HTML内容
        ai_model = data.get('ai_model', 'gpt-4o')  # 选择的AI模型
        
        # 获取模型配置
        model_config = config.AI_MODELS.get(ai_model, config.AI_MODELS['gpt-4o'])
        
        # 执行翻译（使用批量+线程池优化）
        translated_content = translator.translate_batch(
            content, 
            target_lang, 
            source_lang,
            batch_size=config.BATCH_SIZE,
            max_workers=config.MAX_WORKERS,
            model_config=model_config
        )
        
        # 如果有HTML内容，生成翻译后的HTML
        translated_html = None
        if html_content:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 替换每个段落的文本
            for item in translated_content:
                element_id = item.get('id')
                if element_id:
                    element = soup.find(id=element_id)
                    if element:
                        element.string = item.get('translation', '')
            
            translated_html = str(soup)
        
        return jsonify({
            'success': True,
            'translated_content': translated_content,
            'translated_html': translated_html
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/translate-single', methods=['POST'])
def translate_single():
    """翻译单段文本（用于实时翻译）"""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': '缺少文本'}), 400
        
        text = data['text']
        target_lang = data.get('target_lang', 'zh-CN')
        source_lang = data.get('source_lang', 'auto')
        
        # 执行翻译
        translation = translator.translate_text(text, target_lang, source_lang)
        
        return jsonify({
            'success': True,
            'translation': translation
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/translate_image', methods=['POST'])
def translate_image():
    """整图翻译（直接发送图片给AI翻译）"""
    try:
        data = request.get_json()
        
        if not data or 'image_base64' not in data:
            return jsonify({'error': '缺少图片数据'}), 400
        
        image_base64 = data['image_base64']
        target_lang = data.get('target_lang', 'zh-CN')
        ai_model = data.get('ai_model', 'gpt-4o')
        
        # 获取模型配置
        model_config = config.AI_MODELS.get(ai_model, config.AI_MODELS['gpt-4o'])
        
        # 执行整图翻译
        translation = translator.translate_image(
            image_base64,
            target_lang,
            model_config=model_config
        )
        
        return jsonify({
            'success': True,
            'translation': translation
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/export', methods=['POST'])
def export_translation():
    """导出翻译结果（仅译文）"""
    try:
        data = request.get_json()
        
        if not data or 'content' not in data:
            return jsonify({'error': '缺少内容'}), 400
        
        content = data['content']
        format_type = data.get('format', 'txt')  # txt 或 docx
        filename = data.get('filename', '翻译结果')
        has_format = data.get('has_format', False)
        translated_html = data.get('translated_html')
        
        if format_type == 'txt':
            # 导出为TXT文件（仅译文）
            output = io.StringIO()
            output.write(f"翻译结果\n")
            output.write(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            output.write("=" * 80 + "\n\n")
            
            if has_format and translated_html:
                # 从HTML中提取纯文本
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(translated_html, 'html.parser')
                text = soup.get_text(separator='\n')
                output.write(text)
            else:
                # 段落模式，只输出译文
                for item in content:
                    translation = item.get('translation', '')
                    if translation:
                        output.write(translation + "\n\n")
            
            # 转换为字节流
            output_bytes = io.BytesIO()
            output_bytes.write(output.getvalue().encode('utf-8-sig'))  # 使用BOM以支持中文
            output_bytes.seek(0)
            
            return send_file(
                output_bytes,
                mimetype='text/plain',
                as_attachment=True,
                download_name=f'{filename}_译文.txt'
            )
            
        elif format_type == 'docx':
            # 导出为Word文档（仅译文）
            if has_format and translated_html:
                # 使用mammoth将HTML转回Word（保留格式）
                # 注意：mammoth主要用于Word转HTML，这里我们用python-docx从HTML创建
                from bs4 import BeautifulSoup
                doc = Document()
                soup = BeautifulSoup(translated_html, 'html.parser')
                
                # 遍历HTML元素，转换为Word格式
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
                    if element.name.startswith('h'):
                        # 标题
                        level = int(element.name[1])
                        doc.add_heading(element.get_text(), level=level)
                    elif element.name == 'p':
                        # 段落
                        para = doc.add_paragraph(element.get_text())
                        # 保留粗体和斜体
                        if element.find('strong') or element.find('b'):
                            para.runs[0].bold = True
                        if element.find('em') or element.find('i'):
                            para.runs[0].italic = True
                    elif element.name in ['ul', 'ol']:
                        # 列表
                        for li in element.find_all('li'):
                            doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
            else:
                # 段落模式，只导出译文
                doc = Document()
                doc.add_heading('译文', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for item in content:
                    translation = item.get('translation', '')
                    if translation:
                        doc.add_paragraph(translation)
            
            # 保存到内存
            output_bytes = io.BytesIO()
            doc.save(output_bytes)
            output_bytes.seek(0)
            
            return send_file(
                output_bytes,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'{filename}_译文.docx'
            )
        
        else:
            return jsonify({'error': '不支持的导出格式'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=config.DEBUG, host='0.0.0.0', port=5001)
