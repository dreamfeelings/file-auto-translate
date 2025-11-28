"""
文件解析模块 - 支持PDF、Word、图片等文件类型
"""
import os
from typing import List, Dict, Tuple
import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import mammoth
from bs4 import BeautifulSoup
import re
import io
import base64
import requests

class FileParser:
    """文件解析器"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> List[Dict[str, str]]:
        """解析PDF文件（简单段落模式）"""
        content = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        # 按段落分割
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        for para in paragraphs:
                            content.append({
                                'page': page_num,
                                'text': para
                            })
        except Exception as e:
            raise Exception(f"PDF解析错误: {str(e)}")
        
        return content
    
    @staticmethod
    def parse_pdf_with_format(file_path: str) -> Tuple[str, List[Dict[str, str]]]:
        """
        解析PDF文件并保留格式（包括表格）
        返回: (HTML格式的完整文档, 段落列表用于翻译)
        """
        try:
            doc = fitz.open(file_path)
            html_parts = []
            paragraphs = []
            para_index = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # 尝试提取表格
                tables = page.find_tables()
                
                if tables and len(tables.tables) > 0:
                    # 页面包含表格
                    for table in tables:
                        table_data = table.extract()
                        if table_data:
                            # 生成HTML表格
                            html_parts.append('<table class="pdf-table" style="width: 100%; border-collapse: collapse; margin: 20px 0;">')
                            
                            for row_idx, row in enumerate(table_data):
                                html_parts.append('<tr>')
                                for cell_idx, cell in enumerate(row):
                                    # 清理单元格文本
                                    cell_text = str(cell).strip() if cell else ''
                                    
                                    # 修复下划线分离问题：将 "word _ _" 转换为 "word__"
                                    # 移除下划线前后的空格，但保留下划线本身
                                    cell_text = re.sub(r'\s+_\s+', '_', cell_text)  # "word _ word" -> "word_word"
                                    cell_text = re.sub(r'\s+_', '_', cell_text)     # "word _" -> "word_"
                                    cell_text = re.sub(r'_\s+', '_', cell_text)     # "_ word" -> "_word"
                                    
                                    # 只移除尾部的空格，不移除下划线
                                    cell_text = cell_text.rstrip()
                                    
                                    para_id = f'para-{para_index}'
                                    
                                    # 第一行作为表头
                                    if row_idx == 0:
                                        html_parts.append(f'<th id="{para_id}" class="translatable" style="border: 1px solid #ddd; padding: 12px; background: #f5f5f7; text-align: left; font-weight: 600;">{cell_text}</th>')
                                    else:
                                        html_parts.append(f'<td id="{para_id}" class="translatable" style="border: 1px solid #ddd; padding: 12px;">{cell_text}</td>')
                                    
                                    if cell_text:
                                        paragraphs.append({
                                            'id': para_id,
                                            'text': cell_text,
                                            'tag': 'th' if row_idx == 0 else 'td',
                                            'index': para_index,
                                            'page': page_num + 1,
                                            'is_table': True,
                                            'row': row_idx,
                                            'col': cell_idx
                                        })
                                        para_index += 1
                                
                                html_parts.append('</tr>')
                            
                            html_parts.append('</table>')
                else:
                    # 没有表格，按常规文本处理
                    blocks = page.get_text("dict")["blocks"]
                    
                    for block in blocks:
                        if block.get("type") == 0:  # 文本块
                            for line in block.get("lines", []):
                                # 提取每一行的文本
                                line_text = ""
                                font_size = 12
                                is_bold = False
                                
                                for span in line.get("spans", []):
                                    line_text += span.get("text", "")
                                    font_size = span.get("size", 12)
                                    font_flags = span.get("flags", 0)
                                    # 检测是否为粗体 (flag & 16)
                                    is_bold = (font_flags & 16) != 0
                                
                                line_text = line_text.strip()
                                if not line_text:
                                    continue
                                
                                # 过滤掉只包含符号或过短的文本
                                if len(line_text) <= 2 and line_text in ['•', '●', '○', '■', '□', '▪', '▫', '-', '*', '·', '.', ',']:
                                    continue
                                
                                # 过滤掉只包含数字的行（如页码）
                                if line_text.isdigit() and len(line_text) <= 3:
                                    continue
                                
                                # 根据字体大小判断是否为标题
                                para_id = f'para-{para_index}'
                                if font_size > 16:
                                    html_parts.append(f'<h1 id="{para_id}" class="translatable">{line_text}</h1>')
                                    tag = 'h1'
                                elif font_size > 14:
                                    html_parts.append(f'<h2 id="{para_id}" class="translatable">{line_text}</h2>')
                                    tag = 'h2'
                                elif font_size > 13:
                                    html_parts.append(f'<h3 id="{para_id}" class="translatable">{line_text}</h3>')
                                    tag = 'h3'
                                else:
                                    if is_bold:
                                        html_parts.append(f'<p id="{para_id}" class="translatable"><strong>{line_text}</strong></p>')
                                    else:
                                        html_parts.append(f'<p id="{para_id}" class="translatable">{line_text}</p>')
                                    tag = 'p'
                                
                                paragraphs.append({
                                    'id': para_id,
                                    'text': line_text,
                                    'tag': tag,
                                    'index': para_index,
                                    'page': page_num + 1
                                })
                                para_index += 1
                
                # 页面分隔
                if page_num < len(doc) - 1:
                    html_parts.append('<hr style="margin: 20px 0; border: none; border-top: 1px solid #E5E5E5;">')
            
            doc.close()
            
            # 组合HTML
            html_content = '\n'.join(html_parts)
            
            return html_content, paragraphs
            
        except Exception as e:
            raise Exception(f"PDF格式解析错误: {str(e)}")
    
    @staticmethod
    def parse_docx(file_path: str) -> List[Dict[str, str]]:
        """解析Word文档（简单段落模式）"""
        content = []
        try:
            doc = Document(file_path)
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                text = paragraph.text.strip()
                if text:
                    content.append({
                        'paragraph': para_num,
                        'text': text
                    })
        except Exception as e:
            raise Exception(f"Word文档解析错误: {str(e)}")
        
        return content
    
    @staticmethod
    def parse_docx_with_format(file_path: str) -> Tuple[str, List[Dict[str, str]]]:
        """
        解析Word文档并保留格式
        返回: (HTML格式的完整文档, 段落列表用于翻译)
        """
        try:
            # 使用mammoth将Word转换为HTML
            with open(file_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
            
            # 解析HTML提取文本段落
            soup = BeautifulSoup(html_content, 'html.parser')
            paragraphs = []
            
            # 提取所有段落
            for idx, element in enumerate(soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'])):
                text = element.get_text().strip()
                if text:
                    # 为每个元素添加唯一ID，方便后续替换
                    element_id = f'para-{idx}'
                    element['id'] = element_id
                    element['class'] = element.get('class', []) + ['translatable']
                    
                    paragraphs.append({
                        'id': element_id,
                        'text': text,
                        'tag': element.name,
                        'index': idx
                    })
            
            # 返回更新后的HTML和段落列表
            formatted_html = str(soup)
            
            return formatted_html, paragraphs
            
        except Exception as e:
            raise Exception(f"Word文档格式解析错误: {str(e)}")
    
    @staticmethod
    def parse_image(file_path: str, api_key: str = None, api_base_url: str = None, model: str = None, max_retries: int = 3) -> List[Dict[str, str]]:
        """使用AI识别图片文字（支持失败重试）"""
        import config
        import time
        
        content = []
        last_error = None
        
        # 读取图片并转换为base64（只读取一次）
        try:
            with open(file_path, 'rb') as img_file:
                image_data = img_file.read()
                base64_image = base64.b64encode(image_data).decode('utf-8')
            
            # 判断图片格式
            image_ext = file_path.rsplit('.', 1)[1].lower()
            mime_type = f'image/{image_ext if image_ext != "jpg" else "jpeg"}'
        except Exception as e:
            raise Exception(f"图片读取失败: {str(e)}")
        
        # 使用配置中的API信息
        api_key = api_key or config.API_KEY
        api_base_url = api_base_url or config.API_BASE_URL
        model = model or config.MODEL
        
        # 重试机制
        for attempt in range(max_retries):
            try:
                # 调用AI API识别图片
                headers = {
                    'Authorization': f'Bearer {api_key}',
                    'Content-Type': 'application/json'
                }
                
                payload = {
                    'model': model,
                    'messages': [
                        {
                            'role': 'user',
                            'content': [
                                {
                                    'type': 'text',
                                    'text': '请提取图片中的所有文字内容。要求：\n1. 保持原文的段落格式\n2. 不要添加任何解释或注释\n3. 只返回提取的文字内容\n4. 如果有多个段落，用空行分隔'
                                },
                                {
                                    'type': 'image_url',
                                    'image_url': {
                                        'url': f'data:{mime_type};base64,{base64_image}'
                                    }
                                }
                            ]
                        }
                    ],
                    'max_tokens': 4000
                }
                
                response = requests.post(
                    f'{api_base_url}/v1/chat/completions',
                    headers=headers,
                    json=payload,
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    text = result['choices'][0]['message']['content'].strip()
                    
                    if text:
                        # 按段落分割
                        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                        if not paragraphs:
                            # 如果没有空行分隔，按单行分割
                            paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                        
                        for para_num, para in enumerate(paragraphs, 1):
                            if para:
                                content.append({
                                    'paragraph': para_num,
                                    'text': para
                                })
                        
                        # 成功识别，返回结果
                        if attempt > 0:
                            print(f"图片识别成功（第{attempt + 1}次尝试）")
                        return content
                    else:
                        raise Exception("图片中未识别到文字")
                else:
                    raise Exception(f"API返回错误: {response.status_code} - {response.text}")
                    
            except Exception as e:
                last_error = e
                if attempt < max_retries - 1:
                    # 还有重试机会，等待后重试
                    wait_time = (attempt + 1) * 2  # 递增等待时间：2秒、4秒、6秒
                    print(f"图片识别失败（第{attempt + 1}次尝试），{wait_time}秒后重试: {str(e)}")
                    time.sleep(wait_time)
                else:
                    # 已达到最大重试次数
                    print(f"图片识别失败，已重试{max_retries}次")
        
        # 所有重试都失败
        raise Exception(f"图片识别错误（已重试{max_retries}次）: {str(last_error)}")
        
        return content
    
    @staticmethod
    def parse_file(file_path: str, file_type: str) -> List[Dict[str, str]]:
        """统一的文件解析接口"""
        file_ext = file_type.lower()
        
        if file_ext == 'pdf':
            return FileParser.parse_pdf(file_path)
        elif file_ext in ['docx', 'doc']:
            return FileParser.parse_docx(file_path)
        elif file_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
            return FileParser.parse_image(file_path)
        else:
            raise Exception(f"不支持的文件类型: {file_type}")
